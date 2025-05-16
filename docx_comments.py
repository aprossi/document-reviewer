#!/usr/bin/env python3
"""
DOCX Enhanced Feedback Tool with Side Comments
Provides feedback on Word documents using LLMs via Ollama's REST API
Creates a duplicate document with Word comments in the margins
Processes text in paragraphs, tables, and text boxes
"""

from docx import Document
import os
import re
import argparse
import sys
import time
import requests
import json
import docx.shared
import shutil
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

def extract_text_with_locations(doc):
    """
    Extract text with precise location information for better matching
    Args:
        doc: The Word document object
    Returns:
        List of dictionaries with text content and detailed location info
    """
    text_elements = []
    
    # Process main document paragraphs
    print("Extracting text from main document...")
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            text_elements.append({
                'text': para.text,
                'type': 'paragraph',
                'location': {
                    'paragraph_index': i
                },
                'obj': para
            })
    
    # Process tables
    print("Extracting text from tables...")
    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                for para_idx, para in enumerate(cell.paragraphs):
                    if para.text.strip():
                        text_elements.append({
                            'text': para.text,
                            'type': 'table',
                            'location': {
                                'table_index': table_idx,
                                'row_index': row_idx,
                                'cell_index': cell_idx,
                                'paragraph_index': para_idx
                            },
                            'obj': para,
                            'display': f"Table {table_idx+1}, Row {row_idx+1}, Cell {cell_idx+1}, Para {para_idx+1}"
                        })
    
    # Try to find text in shapes/text boxes (more complex)
    print("Attempting to extract text from text boxes and shapes...")
    try:
        # This is more complex but we can try to access shapes
        for i, p in enumerate(doc.paragraphs):
            for run in p.runs:
                if hasattr(run, '_r') and run._r is not None:
                    for drawing in run._r.findall('.//w:drawing', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
                        for txbx in drawing.findall('.//w:txbxContent', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
                            # Text box found, create paragraph objects
                            for para_idx, p_elem in enumerate(txbx.findall('.//w:p', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})):
                                para = Paragraph(p_elem, p._parent)
                                if para.text.strip():
                                    text_elements.append({
                                        'text': para.text,
                                        'type': 'textbox',
                                        'location': {
                                            'paragraph_index': i,
                                            'textbox_paragraph_index': para_idx
                                        },
                                        'obj': para,
                                        'display': f"TextBox in para {i+1}, TextBox Para {para_idx+1}"
                                    })
    except Exception as e:
        print(f"Note: Could not fully process text boxes: {str(e)}")
        print("Text boxes will be excluded from analysis.")
    
    return text_elements

def add_comment_to_paragraph(paragraph, comment_text, author="AI Editor"):
    """
    Add a Word comment to a paragraph using a simpler approach that doesn't require COMMENTS_TYPE.
    
    This uses direct XML manipulation to add comments since python-docx doesn't have built-in
    comment support. It's a simplified approach that works with the current structure.
    
    Args:
        paragraph: The paragraph to comment on
        comment_text: The text of the comment
        author: The comment author name
    """
    # We're going to add a custom attribute to track comment IDs across the document
    if not hasattr(paragraph.part, '_next_comment_id'):
        paragraph.part._next_comment_id = 1
    
    # Get the next comment ID
    comment_id = str(paragraph.part._next_comment_id)
    paragraph.part._next_comment_id += 1

    # Get access to the document's XML tree
    document = paragraph.part.document
    
    # Make sure there's a run in the paragraph (needed for the comment reference)
    if not paragraph.runs:
        paragraph.add_run()
    
    # Get the run where we'll attach the comment reference
    run = paragraph.runs[0]._r  # The XML element of the first run
    
    # Create the comment reference mark
    comment_reference = OxmlElement('w:commentReference')
    comment_reference.set(qn('w:id'), comment_id)
    run.append(comment_reference)
    
    # Define the comments namespace constant 
    comments_type = "{http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments}"
    comments_xml_path = "/word/comments.xml"
    
    # Get access to the package
    package = document.part.package
    
    # Check if there's a comments part already
    comments_part = None
    for rel in document.part.rels.values():
        if rel.reltype == comments_type:
            comments_part = rel.target_part
            break
    
    # If no comments part exists, create it
    if comments_part is None:
        # First, create the Comments XML
        from docx.oxml.parser import parse_xml
        from docx.opc.part import Part
        from docx.opc.packuri import PackURI
        from lxml import etree
        
        # Create a basic empty comments XML structure
        comments_xml = '''
        <w:comments xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
        </w:comments>
        '''
        
        # Create the part manually
        comments_uri = PackURI(comments_xml_path)
        content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml"
        
        # Create a new part with the comments XML
        comments_element = parse_xml(comments_xml)
        
        # Create a proper part with the correct blob property
        class CommentsPart(Part):
            def __init__(self, partname, content_type, element, package):
                super().__init__(partname, content_type, None, package)
                self.element = element
            
            @property
            def blob(self):
                return etree.tostring(self.element, encoding='utf-8', xml_declaration=True)
                
        # Create the part with our custom class that has a proper blob implementation
        comments_part = CommentsPart(comments_uri, content_type, comments_element, package)
        
        # Add the part to the package
        package.parts.append(comments_part)
        
        # Add a relationship from the main document to the comments part
        document.part.rels.add_relationship(comments_type, comments_part, comments_xml_path)
        
        # Store a reference to the comments element for future use
        document.part._comments_element = comments_element
    else:
        # If the comments part exists, get its element
        if not hasattr(document.part, '_comments_element'):
            if hasattr(comments_part, 'element'):
                document.part._comments_element = comments_part.element
            elif hasattr(comments_part, '_element'):
                document.part._comments_element = comments_part._element
            else:
                # Fallback option - try to parse XML from the part's blob
                from docx.oxml.parser import parse_xml
                document.part._comments_element = parse_xml(comments_part.blob)
    
    # Get the comments element
    comments_element = document.part._comments_element
    
    # Create a new comment element
    comment = OxmlElement('w:comment')
    comment.set(qn('w:id'), comment_id)
    comment.set(qn('w:author'), author)
    
    # Get initials (like "AE" for "AI Editor")
    initials = ''.join([word[0].upper() for word in author.split() if word])
    if not initials:
        initials = "AI"
    comment.set(qn('w:initials'), initials)
    
    # Set the date
    comment.set(qn('w:date'), time.strftime('%Y-%m-%dT%H:%M:%SZ'))
    
    # Add the comment text
    p_element = OxmlElement('w:p')
    r_element = OxmlElement('w:r')
    t_element = OxmlElement('w:t')
    t_element.text = comment_text
    r_element.append(t_element)
    p_element.append(r_element)
    comment.append(p_element)
    
    # Add the comment to the comments element
    comments_element.append(comment)
    
    # Make sure the comments part's element attribute is updated
    if hasattr(comments_part, 'element'):
        comments_part.element = comments_element
    elif hasattr(comments_part, '_element'):
        comments_part._element = comments_element

def format_markdown_for_docx(markdown_text, doc):
    """Format markdown text for Word document with proper heading levels"""
    lines = markdown_text.split('\n')
    current_text = ""
    
    for line in lines:
        # Check for headings
        if line.startswith('# '):
            # Add accumulated text before handling new heading
            if current_text:
                doc.add_paragraph(current_text)
                current_text = ""
            
            # Add level 1 heading
            doc.add_heading(line[2:], level=2)
        elif line.startswith('## '):
            # Add accumulated text before handling new heading
            if current_text:
                doc.add_paragraph(current_text)
                current_text = ""
                
            # Add level 2 heading
            doc.add_heading(line[3:], level=3)
        # Check for numbered list items
        elif re.match(r'^\d+\.\s', line):
            # Add accumulated text before handling list item
            if current_text:
                doc.add_paragraph(current_text)
                current_text = ""
                
            # Add the numbered list item
            doc.add_paragraph(line, style='List Number')
        # Regular text
        else:
            if line.strip():  # Non-empty line
                current_text += line + "\n"
            else:  # Empty line - paragraph break
                if current_text:
                    doc.add_paragraph(current_text)
                    current_text = ""
    
    # Add any remaining text
    if current_text:
        doc.add_paragraph(current_text)
    
    return doc

def confirm_ollama_setup(model_name=None, api_host=None):
    """
    Check if Ollama is running and the requested model is available
    Args:
        model_name: Optional specific model to check for
        api_host: Optional API host address (e.g. http://localhost:11434)
    Returns:
        model_name if found, or False if not found
    """
    print("Checking Ollama setup...")
    api_host = api_host or "http://localhost:11434"
    
    try:
        # Try direct API call with requests
        print(f"Using direct API call to: {api_host}")
        try:
            response = requests.get(f"{api_host}/api/tags")
            
            if response.status_code == 200:
                data = response.json()
                available_models = []
                
                for model in data.get("models", []):
                    model_name_str = model.get("name", "")
                    if model_name_str:
                        available_models.append(model_name_str)
                
                if not available_models:
                    print("\n⚠️  No models found in Ollama!")
                    print("Please pull a model first, for example:")
                    print("ollama pull llama3.1")
                    return False
                
                # Print available models for debugging
                print(f"Found {len(available_models)} models: {', '.join(available_models[:5])}...")
                
                # If a specific model is requested, check if it's available
                if model_name:
                    if model_name in available_models:
                        print(f"✅ Found requested model: {model_name}")
                        return model_name
                    else:
                        # Check if the model name might be a partial match (without tags)
                        matching_models = [m for m in available_models if m.startswith(f"{model_name}:") or m == model_name]
                        if matching_models:
                            print(f"✅ Found matching model: {matching_models[0]}")
                            return matching_models[0]
                        
                        print(f"\n⚠️  Requested model '{model_name}' not found in Ollama!")
                        print(f"Available models: {', '.join(available_models)}")
                        return False
                
                # If no specific model requested, prefer Llama models in this order:
                preferred_models = [
                    "llama3.1", "llama3.2", "llama3", 
                    "mistral", "gemma3", "phi3", 
                    "deepseek-r1"
                ]
                
                for preferred in preferred_models:
                    matching = [m for m in available_models if m.startswith(f"{preferred}:") or m == preferred]
                    if matching:
                        print(f"✅ Using model: {matching[0]}")
                        return matching[0]
                
                # If no preferred models, just use the first available
                print(f"✅ Using model: {available_models[0]}")
                return available_models[0]
            else:
                print(f"\n❌ API returned status code {response.status_code}")
                print(f"Response: {response.text}")
                return False
        
        except Exception as e:
            print(f"\n❌ Direct API call failed: {str(e)}")
            return False
    
    except Exception as e:
        print(f"\n❌ Error connecting to Ollama: {str(e)}")
        print("Please make sure Ollama is installed and running.")
        print("Installation instructions: https://ollama.com/download")
        print("\nDiagnostic information:")
        print("- Try running 'ollama list' from the command line")
        print("- Check if Ollama is listening on port 11434")
        print("- If running on a different port, use the --api-host option")
        print("- Check Ollama logs for errors")
        return False

def list_available_models(api_host=None):
    """
    List all available models in Ollama and exit
    Args:
        api_host: Optional API host address
    """
    api_host = api_host or "http://localhost:11434"
    
    try:
        # Try direct API call with requests
        print(f"Connecting to Ollama API at: {api_host}")
        
        try:
            response = requests.get(f"{api_host}/api/tags")
            
            if response.status_code == 200:
                data = response.json()
                available_models = data.get("models", [])
                
                # Sort models by name
                available_models.sort(key=lambda x: x.get("name", ""))
                
                # Print in a table format
                print("\n🤖 Available Models in Ollama:")
                print("-" * 80)
                print(f"{'MODEL NAME':<40} {'SIZE':<10} {'MODIFIED':<20}")
                print("-" * 80)
                
                for model in available_models:
                    model_name = model.get("name", "")
                    
                    # Format the size in a human-readable format
                    size_bytes = model.get("size", 0)
                    size_gb = size_bytes / 1_000_000_000
                    size_str = f"{size_gb:.1f} GB"
                    
                    # Format the modified date
                    modified = model.get("modified_at", "")
                    if "T" in modified:
                        modified = modified.split("T")[0]
                    
                    print(f"{model_name:<40} {size_str:<10} {modified:<20}")
                
                print("-" * 80)
                print(f"Total models: {len(available_models)}")
                print(f"\nOllama API host: {api_host}")
                print("\nUse --model MODEL_NAME to select a specific model for analysis.")
                print("Example: python ollama-docx-comments-enhanced.py document.docx --model mistral")
                
            else:
                print(f"\n❌ API returned status code {response.status_code}")
                print(f"Response: {response.text}")
        
        except Exception as e:
            print(f"\n❌ Error connecting to Ollama API: {str(e)}")
            print(f"Tried to connect to: {api_host}")
            print("Please make sure Ollama is installed and running.")
    
    except Exception as e:
        print(f"\n❌ Error: {str(e)}")
    
    sys.exit(0)

def analyze_document_with_comments(docx_path, model_name="llama3.1", verbose=False, 
                                 temperature=0.2, api_host=None, system_prompt=None,
                                 summary_doc=True):
    """
    Process a document with LLM and add side comments, handling tables and textboxes
    Args:
        docx_path: Path to the DOCX file
        model_name: Ollama model name to use
        verbose: Whether to print detailed progress
        temperature: Model temperature (higher = more creative)
        api_host: Optional API host address
        system_prompt: Custom system prompt to use (if None, uses default)
        summary_doc: Whether to create a separate summary document
    Returns:
        Path to the created commented document and summary document (if created)
    """
    api_host = api_host or "http://localhost:11434"
    
    try:
        # Read the document
        print(f"Reading document: {docx_path}")
        doc = Document(docx_path)
        
        # Create a duplicate of the original document for comments
        base_name = os.path.basename(docx_path)
        name_without_ext = os.path.splitext(base_name)[0]
        model_short_name = model_name.split(':')[0].replace("/", "-")
        
        # Create paths for the output files
        commented_path = os.path.join(os.path.dirname(docx_path), 
                                    f"{name_without_ext}_comments_{model_short_name}.docx")
        
        # IMPORTANT: Delete the file if it exists to prevent duplicate XML issues
        if os.path.exists(commented_path):
            try:
                os.remove(commented_path)
                print(f"Removed existing file: {commented_path}")
            except Exception as e:
                print(f"Warning: Could not remove existing file: {e}")
        
        # Create a duplicate of the original document
        print(f"Creating a duplicate document for comments...")
        shutil.copy2(docx_path, commented_path)
        
        # Open the duplicate document for editing
        commented_doc = Document(commented_path)
        
        # Create a summary document if requested
        summary_doc_obj = None
        if summary_doc:
            summary_doc_obj = Document()
            summary_doc_obj.add_heading(f"Document Feedback Summary", 0)
            summary_doc_obj.add_paragraph(f"Document: {os.path.basename(docx_path)}")
            summary_doc_obj.add_paragraph(f"Analysis by: {model_name}")
            summary_doc_obj.add_paragraph(f"Date: {time.strftime('%Y-%m-%d %H:%M:%S')}")
        
        # Extract all text with detailed location information
        text_elements = extract_text_with_locations(doc)
        
        # Get text for full document analysis
        all_text = [elem['text'] for elem in text_elements]
        full_text = "\n\n".join(all_text)
        char_count = len(full_text)
        
        # Count by type
        main_paragraphs = sum(1 for elem in text_elements if elem['type'] == 'paragraph')
        table_cells = sum(1 for elem in text_elements if elem['type'] == 'table')
        text_boxes = sum(1 for elem in text_elements if elem['type'] == 'textbox')
        
        print(f"Document contains {char_count} characters across {len(text_elements)} text elements")
        print(f"  - Main paragraphs: {main_paragraphs}")
        print(f"  - Table cells: {table_cells}")
        print(f"  - Text boxes: {text_boxes}")
        
        if char_count == 0:
            print("❌ Document appears to be empty. Please check the file.")
            return None
        
        # Use default system prompt if none provided
        if system_prompt is None:
            system_prompt = '''You are a professional editor providing detailed stylistic feedback.
            Analyze the document for:
            1. Writing style and tone
            2. Clarity and conciseness
            3. Sentence structure and flow
            4. Word choice and vocabulary
            5. Overall impressions and suggestions
            
            Format your feedback as follows:
            
            # Overall Assessment
            [General feedback about the entire document]
            
            # Specific Issues
            1. [First specific issue - be precise]
            2. [Second specific issue]
            [etc.]
            
            # Recommended Improvements
            1. [First recommendation]
            2. [Second recommendation]
            [etc.]
            '''
        
        # Get overall stylistic feedback from the LLM using direct API calls
        print(f"Getting overall document assessment using {model_name}...")
        print(f"Temperature: {temperature:.1f}" + (" (creative mode)" if temperature > 0.5 else ""))
        start_time = time.time()
        
        # Prepare messages
        messages = [
            {
                'role': 'system',
                'content': system_prompt
            },
            {
                'role': 'user',
                'content': f"Please provide stylistic feedback on this document:\n\n{full_text}"
            }
        ]
        
        # Make API call
        try:
            api_data = {
                "model": model_name,
                "messages": messages,
                "options": {
                    "temperature": temperature
                },
                "stream": False
            }
            
            response = requests.post(
                f"{api_host}/api/chat",
                json=api_data
            )
            
            if response.status_code != 200:
                print(f"\n❌ API returned status code {response.status_code}")
                print(f"Response: {response.text}")
                return None
            
            response_data = response.json()
            overall_feedback = response_data.get("message", {}).get("content", "")
            
            if not overall_feedback:
                print("\n❌ No feedback received from model")
                return None
            
        except Exception as e:
            print(f"\n❌ Error calling Ollama API: {str(e)}")
            return None
        
        analysis_time = time.time() - start_time
        print(f"Overall assessment completed in {analysis_time:.1f} seconds")
        
        # Add overall feedback to the summary document
        if summary_doc_obj:
            summary_doc_obj.add_heading("Overall Assessment", 1)
            format_markdown_for_docx(overall_feedback, summary_doc_obj)
            summary_doc_obj.add_heading("Content-by-Content Feedback", 1)
        
        # Count text elements that will be analyzed
        substantive_text = [elem for elem in text_elements if len(elem['text'].strip()) > 30]
        total_to_analyze = len(substantive_text)
        
        if total_to_analyze > 20:
            print(f"⚠️  Document has {total_to_analyze} substantive text elements.")
            print("   Analysis may take some time.")
            
            if total_to_analyze > 50:
                while True:
                    response = input("Document is very large. Continue? (y/n): ").lower()
                    if response in ('y', 'yes'):
                        break
                    elif response in ('n', 'no'):
                        print("Operation cancelled.")
                        return None
        
        # Process each text element
        analyzed_count = 0
        feedback_count = 0
        comment_errors = 0
        print("Analyzing individual text elements and adding comments...")
        
        # Set to track which paragraphs we've already processed
        # This helps when identical text appears multiple times
        processed_elements = set()
        
        # Process each text element
        for i, element in enumerate(text_elements):
            text_content = element['text']
            element_type = element['type']
            location = element['location']
            
            # Skip already processed elements or empty text
            element_key = f"{i}:{text_content[:50]}"
            if element_key in processed_elements or not text_content.strip():
                continue
            
            processed_elements.add(element_key)
            
            # Get a display string for the source
            if element_type == 'paragraph':
                source_display = f"Paragraph {location['paragraph_index']+1}"
            elif element_type == 'table':
                source_display = element.get('display') or f"Table content"
            elif element_type == 'textbox':
                source_display = element.get('display') or f"Text box"
            else:
                source_display = f"Element {i+1}"
            
            # Find the corresponding paragraph in the target document
            target_para = None
            
            # For main paragraphs, use direct index
            if element_type == 'paragraph':
                para_idx = location['paragraph_index']
                if para_idx < len(commented_doc.paragraphs):
                    target_para = commented_doc.paragraphs[para_idx]
                    # Verify content
                    if target_para.text != text_content:
                        target_para = None
            
            # For table content, navigate to the specific table cell
            elif element_type == 'table':
                try:
                    table_idx = location['table_index']
                    row_idx = location['row_index']
                    cell_idx = location['cell_index']
                    para_idx = location['paragraph_index']
                    
                    if table_idx < len(commented_doc.tables):
                        table = commented_doc.tables[table_idx]
                        if row_idx < len(table.rows):
                            row = table.rows[row_idx]
                            if cell_idx < len(row.cells):
                                cell = row.cells[cell_idx]
                                if para_idx < len(cell.paragraphs):
                                    target_para = cell.paragraphs[para_idx]
                                    # Verify content
                                    if target_para.text != text_content:
                                        target_para = None
                except (KeyError, IndexError) as e:
                    if verbose:
                        print(f"\nError accessing table element: {e}")
            
            # If the target paragraph wasn't found, search by content
            if target_para is None:
                # First look in document paragraphs
                for para in commented_doc.paragraphs:
                    if para.text == text_content:
                        target_para = para
                        break
                
                # If not found, look in tables
                if target_para is None:
                    for table in commented_doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                for para in cell.paragraphs:
                                    if para.text == text_content:
                                        target_para = para
                                        break
                                if target_para:
                                    break
                            if target_para:
                                break
                        if target_para:
                            break
            
            # Only analyze substantive text if we found the target
            if len(text_content.strip()) > 30 and target_para is not None:
                analyzed_count += 1
                progress = f"[{analyzed_count}/{total_to_analyze}]"
                
                if verbose:
                    print(f"{progress} Analyzing {element_type} {source_display}...")
                else:
                    # Print progress bar
                    progress_pct = analyzed_count*100//total_to_analyze
                    progress_bar = "█" * (progress_pct//5) + "░" * (20 - progress_pct//5)
                    sys.stdout.write(f"\r{progress} Analyzing content... {progress_bar} {progress_pct}%")
                    sys.stdout.flush()
                
                # Prepare messages for text analysis
                content_messages = [
                    {
                        'role': 'system',
                        'content': '''You are a professional writing coach providing targeted feedback.
                        Review the text and identify any grammar, style, or clarity issues.
                        If you find issues, provide ONE clear, specific suggestion that would most improve it.
                        Be precise and helpful. Limit your feedback to 1-2 sentences.
                        If the text has no significant issues, respond with "No issues found."'''
                    },
                    {
                        'role': 'user',
                        'content': f"Review this text (from {element_type} {source_display}):\n\n{text_content}"
                    }
                ]
                
                # Get feedback for this text element using direct API
                try:
                    api_data = {
                        "model": model_name,
                        "messages": content_messages,
                        "options": {
                            "temperature": temperature
                        },
                        "stream": False
                    }
                    
                    response = requests.post(
                        f"{api_host}/api/chat",
                        json=api_data
                    )
                    
                    if response.status_code != 200:
                        raise RuntimeError(f"API returned status code {response.status_code}")
                    
                    response_data = response.json()
                    content_feedback = response_data.get("message", {}).get("content", "").strip()
                    
                    # Add feedback if there's something to improve
                    if content_feedback and "No issues found" not in content_feedback:
                        feedback_count += 1
                        
                        try:
                            # Add comment to the paragraph
                            add_comment_to_paragraph(target_para, content_feedback)
                            
                            # Add to summary doc if requested
                            if summary_doc_obj:
                                # Add text excerpt
                                excerpt = text_content[:100] + "..." if len(text_content) > 100 else text_content
                                summary_para = summary_doc_obj.add_paragraph()
                                summary_para.add_run(f"{element_type} {source_display}: ").bold = True
                                summary_para.add_run(f"{excerpt}")
                                
                                # Add feedback
                                feedback_para = summary_doc_obj.add_paragraph()
                                feedback_para.style = 'Quote'
                                feedback_para.add_run(f"✎ {content_feedback}").italic = True
                                
                                # Add space after feedback
                                summary_doc_obj.add_paragraph("")
                        except Exception as comment_error:
                            comment_errors += 1
                            if verbose or comment_errors <= 3:  # Only show the first few errors
                                print(f"\nError adding comment to {element_type} {source_display}: {str(comment_error)}")
                            elif comment_errors == 4:
                                print("\nMultiple comment errors occurred. Suppressing further error messages...")
                        
                except Exception as para_e:
                    if verbose:
                        print(f"\nError analyzing {element_type} {source_display}: {str(para_e)}")
            elif target_para is None and len(text_content.strip()) > 30:
                # Element couldn't be found in the target document
                if verbose:
                    print(f"\nCouldn't locate target for {element_type} content: {text_content[:50]}...")
        
        # Clear progress bar line
        if not verbose:
            sys.stdout.write("\r" + " " * 80 + "\r")
            sys.stdout.flush()
        
        # Save the commented document
        commented_doc.save(commented_path)
        
        # Save summary document if created
        summary_path = None
        if summary_doc_obj:
            summary_path = os.path.join(os.path.dirname(docx_path),
                                      f"{name_without_ext}_summary_{model_short_name}.docx")
            summary_doc_obj.save(summary_path)
        
        print(f"\n✅ Document analyzed with {feedback_count} comments added")
        if comment_errors > 0:
            print(f"⚠️  {comment_errors} comments could not be added due to technical limitations")
            print("    The summary document contains all feedback items.")
        
        print(f"✅ Commented document created: {commented_path}")
        if summary_path:
            print(f"✅ Summary document created: {summary_path}")
        print(f"\nNote: For best viewing of comments, open the document in Microsoft Word.")
        
        return commented_path, summary_path
        
    except Exception as e:
        print(f"\n❌ Error: {str(e)}")
        import traceback
        traceback.print_exc()
        return None

def main():
    """Main function to run the script"""
    parser = argparse.ArgumentParser(
        description="Add side comments to a DOCX document using LLMs with Ollama",
        formatter_class=argparse.ArgumentDefaultsHelpFormatter
    )
    parser.add_argument("docx_path", help="Path to the DOCX file to analyze", nargs="?")
    parser.add_argument("--model", "-m", help="Specific model to use (e.g., llama3.1, mistral, phi3)", default=None)
    parser.add_argument("--verbose", "-v", action="store_true", help="Print detailed progress")
    parser.add_argument("--list-models", "-l", action="store_true", help="List available models and exit")
    parser.add_argument("--system-prompt", "-s", help="Path to a custom system prompt file")
    parser.add_argument("--creative", "-c", action="store_true", 
                       help="Use a higher temperature for more creative feedback")
    parser.add_argument("--api-host", help="Ollama API host (e.g., http://localhost:11434)", 
                       default="http://localhost:11434")
    parser.add_argument("--no-summary", action="store_true", 
                       help="Don't create a separate summary document")
    parser.add_argument("--version", action="version", version="DOCX Comments Feedback Tool v1.0.0")
    args = parser.parse_args()
    
    # Handle the list-models argument
    if args.list_models:
        list_available_models(api_host=args.api_host)
        return
    
    if not args.docx_path:
        if not args.list_models:
            parser.print_help()
        return
        
    if not os.path.exists(args.docx_path):
        print(f"❌ Error: File not found - {args.docx_path}")
        return
    
    if not args.docx_path.lower().endswith(".docx"):
        print(f"❌ Error: File must be a .docx file - {args.docx_path}")
        return
    
    # Check Ollama setup with the specified model if provided
    model_name = confirm_ollama_setup(args.model, api_host=args.api_host)
    if not model_name:
        return
    
    # Check if custom system prompt is provided
    system_prompt = None
    if args.system_prompt:
        if os.path.exists(args.system_prompt):
            with open(args.system_prompt, 'r') as f:
                system_prompt = f.read()
            print(f"Using custom system prompt from: {args.system_prompt}")
        else:
            print(f"⚠️ Warning: System prompt file not found: {args.system_prompt}")
            print("Using default system prompt instead.")
    
    # Analyze the document with side comments
    analyze_document_with_comments(
        docx_path=args.docx_path, 
        model_name=model_name,
        verbose=args.verbose,
        temperature=0.7 if args.creative else 0.2,
        api_host=args.api_host,
        system_prompt=system_prompt,
        summary_doc=not args.no_summary
    )

if __name__ == "__main__":
    main()